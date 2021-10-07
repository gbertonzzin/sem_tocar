"""
!!!UNTESTED!!!
Handles invite creating and printing

"""

import os
import win32api
import win32print
from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from modules.sem_tocar_config import *
#import pycups


def write_invitation(event, qr_file):
    """
    Creates a nice .docx invitation for printing
    """
    company = COMPANY_NAME
    contact_phone = CONTACT_PHONE
    host = "Dr. Fulaninho"          #Maybe use the calendar name?
    time = "a hora H"               #Extract from event object
    date = "D"
    
    document = Document()
    
    document.add_paragraph().add_run(f"Sua visita em {company} com {host} está marcada para {time} no dia {date}")
    document.add_paragraph().add_run(f"Ao chegar, apresente o QR abaixo à câmera da entrada.")
    document.add_paragraph().add_run(f"Em caso de problemas, ligue {contact_phone}")
    
    for par in document.paragraphs: #Hope this works
        par.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
    for run in document.runs:       #Hope this works too
        run.add_break()
        run.font.bold = True
    
    document.add_picture(qr_file, width = Cm(10))
    
    document.save(f"docx//{host}_{event['id']}.docx")        #Better save these, right?



def win_print_doc(file): #from http://timgolden.me.uk/python/win32_how_do_i/print.html , works like a charm
    """
    Prints a file in a Windows environment, using the default printer
    """
    win32api.ShellExecute (0,"print",file,f"/d:{win32print.GetDefaultPrinter()}",".",0)


def linux_print_doc(file): #TODO
    """
    Prints a file in a Linux environment
    """
    os.system("lpr -P {DEFAULT_PRINTER} {file}") #Yet untested